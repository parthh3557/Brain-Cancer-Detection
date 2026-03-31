import os
import subprocess
import sys

# Ensure python-docx is installed
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Mid-Point Project Progress Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Project Title
    add_heading(doc, '1. Project Title')
    add_paragraph(doc, 'Brain Cancer Detection Application using ResNet-inspired Custom CNN and Flask deployed on GCP.')

    # 2. Project Team Details
    add_heading(doc, '2. Project Team Details')
    add_paragraph(doc, 'Student Names: [Insert Names Here]\nRoll Numbers: [Insert Roll Numbers Here]')

    # 3. Project Abstract
    add_heading(doc, '3. Project Abstract')
    add_paragraph(doc, 'The project aims to develop an automated Brain Cancer classification web application using deep learning. We developed a custom ResNet-inspired Convolutional Neural Network (CNN) in PyTorch to classify MRI scans into three categories: Glioma, Meningioma, and general Tumor. The model is integrated into a Flask backend, featuring a responsive front-end. To enhance accessibility, we integrated the Google Gemini API to generate easily understandable summaries for laypeople regarding the detected tumor types. The entire system is containerized with Docker and deployed securely on the Google Cloud Platform (Cloud Run) utilizing CPU-optimized PyTorch and Gunicorn for scalable and fast inference.')

    # 4. Problem Statement
    add_heading(doc, '4. Problem Statement')
    add_paragraph(doc, 'Brain cancer diagnosis from MRI scans is complex, time-consuming, and highly dependent on radiologist expertise. Misdiagnosis or delayed diagnosis can severely impact patient survival rates. There is a critical need for an automated, highly accurate, and accessible decision-support system that can assist medical professionals in rapidly identifying tumor types from scans, while also providing patient-friendly explanations for better doctor-patient communication.')

    # 5. Field / Area of the Project
    add_heading(doc, '5. Field / Area of the Project')
    add_paragraph(doc, '• Artificial Intelligence / Deep Learning\n• Computer Vision / Medical Imaging\n• Application Domain: Healthcare / Oncology')

    # 6. Literature Review and Prior Work
    add_heading(doc, '6. Literature Review and Prior Work')
    add_heading(doc, '6.1 Existing Technologies', level=2)
    add_paragraph(doc, 'Existing diagnostic workflows depend entirely on manual MRI review. While some recent research utilizes standard architectures like VGG-16 or traditional Machine Learning, they often suffer from high computational overhead, vanishing gradients in deep networks, or lack the latency optimization needed for web implementation.')
    
    add_heading(doc, '6.2 Related Patents / Research Papers', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Publication'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Limitations'
    hdr_cells[3].text = 'Relevance'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Deep Residual Learning for Image Recognition'
    row_cells[1].text = 'Introduces ResNet architecture utilizing skip connections to solve the vanishing gradient problem in deep CNNs.'
    row_cells[2].text = 'Extremely high computational cost for inference if unmodified.'
    row_cells[3].text = 'Directly inspired our custom lightweight BrainNet50 (1-2-3 custom block) architecture.'

    # 7. Identified Technical Gap
    add_heading(doc, '7. Identified Technical Gap')
    add_paragraph(doc, '• High latency in conventional, manual diagnosis processes.\n• Many existing state-of-the-art models are too heavy for rapid web inference.\n• Lack of an integrated platform combining scientific classification with contextual, patient-friendly education.\n• Sub-optimal deployment configurations resulting in poor model scalability.')

    # 8. Objectives of the Project
    add_heading(doc, '8. Objectives of the Project')
    add_paragraph(doc, 'Primary Objectives:\n• Develop a custom DL-based model (ResNet variant) for robust brain MRI classification.\n• Build a high-performance inference pipeline using a robust Flask application.')
    add_paragraph(doc, 'Secondary Objectives:\n• Containerize the application using Docker focusing on lean CPU-based inference.\n• Explore accessible cloud deployment options via GCP (Google Cloud Platform).\n• Develop a user-friendly frontend prototype integrating LLM summaries via the Gemini API.')

    # 9. Proposed System Overview
    add_heading(doc, '9. Proposed System Overview')
    add_paragraph(doc, 'The system features a user-friendly web interface where an MRI image is uploaded. The Flask backend applies essential transformations such as resizing and normalization before passing the image tensor into a custom PyTorch CNN model. The neural network infers the tumor class alongside a confidence score. This result is immediately returned to the frontend. Furthermore, users can click a button to retrieve a non-medical layman’s summary generated securely via the Gemini API, with the entirety of the architecture hosted robustly on GCP.')

    # 10. Preliminary System Architecture
    add_heading(doc, '10. Preliminary System Architecture')
    add_paragraph(doc, '• Data Input Layer: HTML/JS Interface for MRI image upload via HTTP POST request.\n• Data Processing Layer: torchvision.transforms (Resize to 224x224, Convert to Tensor, Normalize with specific mean/std).\n• Deep Learning Model: BrainNet50 (Custom PyTorch implementation with residual blocks, utilizing Conv2d, BatchNorm2d, and ReLU).\n• Decision Logic: torch.argmax applied to softmax probabilities for classification; integration with Gemini API system prompting.\n• Output / User Interface: Predicted Class, Confidence Percentage, and AI-generated descriptive summary.')

    # 11. Dataset Planning
    add_heading(doc, '11. Dataset Planning')
    add_paragraph(doc, 'Dataset Source: Medical Imaging Dataset (e.g., Kaggle Brain Tumor Classification MRIs)\nExpected Dataset Size: Thousands of high-resolution MRI scans.\nData Collection Method: Compilation and sanitization of publicly available, anonymized clinical datasets.\nPreprocessing Plan: Resizing all inputs to (224, 224), transforming images into tensors, and normalizing using ImageNet statistics (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]).')

    # 12. Model Selection Plan
    add_heading(doc, '12. Model Selection Plan')
    add_paragraph(doc, 'Proposed Model Type: Custom ResNet-inspired Architecture (Deep CNN with Residual Connections)\nTraining Framework: PyTorch\nInitial Architecture Idea: We implemented `BrainCancerModelV1`, consisting of heavily customized repeating `block` layers. Each residual block employs three convolutional layers (1x1, 3x3, 1x1 expansions) ensuring feature retention without gradient degradation. The network is finalized via an AdaptiveAvgPool2d layer and a Fully Connected linear map for the specific classes: Glioma, Meningioma, and general Tumor.')

    # 13. Work Completed So Far
    add_heading(doc, '13. Work Completed So Far')
    add_paragraph(doc, '• Conducted literature review focusing on residual networks for medical imaging.\n• Developed, trained, and saved the custom PyTorch ResNet model (`BrainCancerModel1_ver2.pth`).\n• Developed the Flask backend incorporating endpoints for computer vision inference (`/predict`) and LLM summary generation (`/gemini_summary`).\n• Designed and integrated the UI frontend (`test.html`).\n• Prepared the application for production deployment utilizing a custom `Dockerfile` prioritizing CPU-optimization (`python:3.10-slim` image) and a `gunicorn` WSGI server.\n• Successful deployment to the Google Cloud Platform via Cloud Run.')

    # 14. Preliminary Results
    add_heading(doc, '14. Preliminary Results')
    add_paragraph(doc, 'The model successfully processes user-uploaded MRI scans and returns classifications across the three specified categories. By optimizing the Torch ecosystem within the Dockerfile, we achieved significantly fast CPU-based inference latency. The API bindings with Google\'s Gemini 2.5 Flash API function correctly, returning concise, structured summaries without exposing backend credentials or yielding unprompted medical advice.')

    # 15. Work Plan for the Remaining Month
    add_heading(doc, '15. Work Plan for the Remaining Month')
    add_paragraph(doc, 'Week 1: Finalize frontend styling and implement polished loading/transition states.\nWeek 2: Conduct model generalization and robustness tests over unseen evaluation distributions.\nWeek 3: Refine the Gemini system prompts with rigorous adversarial testing ensuring absolute compliance with non-medical advice constraints.\nWeek 4: Formulate final project documentation, User Guides, and presentation materials.')

    # 16. Expected Novelty / Patent Potential
    add_heading(doc, '16. Expected Novelty / Patent Potential')
    add_paragraph(doc, 'The project\'s core novelty resides in bridging deterministic deep learning predictions with non-deterministic Large Language Model insights within a single interface tailored for patient and clinical education. Furthermore, our customized, lightweight ResNet-variant architecture carefully balances inference speed with accuracy, overcoming the traditional limits of heavy medical CNNs in continuous web deployments.')

    # 17. Challenges Faced So Far
    add_heading(doc, '17. Challenges Faced So Far')
    add_paragraph(doc, 'A major challenge was ensuring the deep learning model footprint remained small enough for cost-effective cloud deployment. This was overcome by writing a specific `Dockerfile` that ensures only the CPU version of PyTorch is installed, saving substantial container storage space. Additionally, securely integrating the Google Gemini API key into our GCP deployment flow without exposing secrets in our codebase required implementing environment variables (via `os.getenv` and `.env` setups).')

    doc.save('Progress_Report.docx')
    print("Report generated successfully: Progress_Report.docx")

if __name__ == '__main__':
    create_report()
