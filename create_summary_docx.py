import json
import docx
from docx.shared import Pt
import sys

notebook_path = 'd:/brain_flask_app/braincancerdetection-vit.ipynb'

with open(notebook_path, 'r', encoding='utf-8') as f:
    nb = json.load(f)

summary_text = None

# Search through cells for the summary output
for i, cell in enumerate(nb.get('cells', [])):
    if cell.get('cell_type') == 'code':
        # Check if output contains a summary
        for output in cell.get('outputs', []):
            if output.get('output_type') == 'stream' and output.get('name') == 'stdout':
                text = "".join(output.get('text', []))
                if 'Total params' in text and ('Layer' in text or 'Conv2d' in text or 'Linear' in text):
                    summary_text = text
                    print(f"Found summary output stream in cell index {i}, execution count {cell.get('execution_count')}")
                    break
            
            elif output.get('output_type') == 'execute_result':
                data = output.get('data', {})
                if 'text/plain' in data:
                    text = "".join(data['text/plain'])
                    if 'Total params' in text and ('Layer' in text or 'Conv2d' in text or 'Linear' in text):
                        summary_text = text
                        print(f"Found summary execute_result in cell index {i}, execution count {cell.get('execution_count')}")
                        break
        
        if summary_text:
            break

if not summary_text:
    # Let's try to find an exact match if not found by heuristics
    # Cell 27 might refer to index 26
    if len(nb.get('cells', [])) > 26:
        cell_27 = nb['cells'][26]
        if cell_27.get('cell_type') == 'code':
            for output in cell_27.get('outputs', []):
                if output.get('output_type') == 'stream':
                    summary_text = "".join(output.get('text', []))
                    print("Extracted from index 26 cell outputs.")
                    break

if summary_text:
    print("Output length:", len(summary_text))
    doc = docx.Document()
    doc.add_heading('Model Architecture Summary', 0)
    p = doc.add_paragraph(summary_text)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    # Let's clean up any weird newlines
    doc.save('d:/brain_flask_app/model_architecture.docx')
    print("Created model_architecture.docx successfully.")
else:
    print("Could not find the model summary output in the notebook.", file=sys.stderr)
